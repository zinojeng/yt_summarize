#!/usr/bin/env python3

"""
驗證 DOCX 文件中的格式轉換是否正確
"""

from docx import Document
import sys

def verify_docx_formatting(docx_path):
    """驗證 DOCX 文件的格式"""
    try:
        doc = Document(docx_path)
        
        print(f"📄 分析文件: {docx_path}")
        print(f"📊 段落總數: {len(doc.paragraphs)}")
        
        bold_runs = 0
        italic_runs = 0
        normal_runs = 0
        code_runs = 0
        
        # 檢查前20個段落的格式
        for i, para in enumerate(doc.paragraphs[:20]):
            if not para.text.strip():
                continue
                
            print(f"\n段落 {i+1}: {para.text[:100]}...")
            
            for j, run in enumerate(para.runs):
                if not run.text.strip():
                    continue
                    
                formats = []
                if run.bold:
                    formats.append("粗體")
                    bold_runs += 1
                if run.italic:
                    formats.append("斜體")
                    italic_runs += 1
                if run.font.name == 'Consolas':
                    formats.append("代碼")
                    code_runs += 1
                
                if not formats:
                    formats.append("普通")
                    normal_runs += 1
                
                format_str = ", ".join(formats)
                print(f"  Run {j+1}: '{run.text[:50]}...' -> [{format_str}]")
        
        print(f"\n📈 格式統計:")
        print(f"  粗體文字段數: {bold_runs}")
        print(f"  斜體文字段數: {italic_runs}")
        print(f"  代碼文字段數: {code_runs}")
        print(f"  普通文字段數: {normal_runs}")
        
        if bold_runs > 0:
            print("✅ 發現粗體格式 - 轉換成功！")
        else:
            print("❌ 未發現粗體格式 - 可能轉換失敗")
            
        return bold_runs > 0
        
    except Exception as e:
        print(f"❌ 分析失敗: {e}")
        return False

if __name__ == "__main__":
    docx_files = [
        "fixed_test.docx",
        "improved_test_f_ui2grDdRMwqOy-5yix0g.docx"
    ]
    
    for docx_file in docx_files:
        try:
            success = verify_docx_formatting(docx_file)
            print(f"{'✅' if success else '❌'} {docx_file}")
            print("=" * 60)
        except FileNotFoundError:
            print(f"⚠️ 文件不存在: {docx_file}")
            print("=" * 60)