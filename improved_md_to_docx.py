#!/usr/bin/env python3

"""
Improved Markdown to DOCX Converter using markdown-it-py
基於 markdown-it-py 的改進版 Markdown 轉 DOCX 轉換器
"""

import logging
from typing import Dict, List, Optional, Any
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown_it import MarkdownIt
from markdown_it.token import Token

logger = logging.getLogger(__name__)


class ImprovedMarkdownToDocxConverter:
    """改進的 Markdown 轉 DOCX 轉換器"""
    
    def __init__(self):
        self.doc = Document()
        self.md = MarkdownIt()
        self.setup_styles()
    
    def setup_styles(self):
        """設置 Word 文檔樣式"""
        try:
            # 設置正文樣式
            normal_style = self.doc.styles['Normal']
            normal_font = normal_style.font
            normal_font.name = '微軟正黑體'
            normal_font.size = Pt(11)
            
            # 創建代碼塊樣式
            try:
                code_style = self.doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
                code_font = code_style.font
                code_font.name = 'Consolas'
                code_font.size = Pt(10)
                code_style.paragraph_format.left_indent = Inches(0.5)
                code_style.paragraph_format.space_before = Pt(6)
                code_style.paragraph_format.space_after = Pt(6)
            except:
                pass  # 樣式可能已存在
            
        except Exception as e:
            logger.warning(f"設置樣式時發生錯誤: {e}")
    
    def convert_markdown_to_docx(self, markdown_content: str, title: str = "文檔") -> BytesIO:
        """
        將 Markdown 內容轉換為 DOCX 文檔
        
        Args:
            markdown_content (str): Markdown 格式的內容
            title (str): 文檔標題
            
        Returns:
            BytesIO: DOCX 文檔的二進制流
        """
        try:
            # 添加文檔標題
            if title and title != "文檔":
                title_para = self.doc.add_heading(title, level=0)
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 使用 markdown-it-py 解析 Markdown
            tokens = self.md.parse(markdown_content)
            
            # 處理解析後的 tokens
            self._process_tokens(tokens)
            
            # 生成 DOCX 二進制流
            docx_stream = BytesIO()
            self.doc.save(docx_stream)
            docx_stream.seek(0)
            
            return docx_stream
            
        except Exception as e:
            logger.error(f"轉換 Markdown 到 DOCX 時發生錯誤: {e}")
            raise
    
    def _process_tokens(self, tokens: List[Token]):
        """處理解析後的 tokens"""
        i = 0
        while i < len(tokens):
            token = tokens[i]
            
            if token.type == 'heading_open':
                i = self._process_heading(tokens, i)
            elif token.type == 'paragraph_open':
                i = self._process_paragraph(tokens, i)
            elif token.type == 'bullet_list_open':
                i = self._process_list(tokens, i, ordered=False)
            elif token.type == 'ordered_list_open':
                i = self._process_list(tokens, i, ordered=True)
            elif token.type == 'blockquote_open':
                i = self._process_blockquote(tokens, i)
            elif token.type == 'fence' or token.type == 'code_block':
                self._process_code_block(token)
                i += 1
            elif token.type == 'hr':
                self._add_horizontal_line()
                i += 1
            else:
                i += 1
    
    def _process_heading(self, tokens: List[Token], start_idx: int) -> int:
        """處理標題"""
        open_token = tokens[start_idx]
        level = int(open_token.tag[1])  # h1 -> 1, h2 -> 2, etc.
        
        # 找到標題內容
        content_idx = start_idx + 1
        if content_idx < len(tokens) and tokens[content_idx].type == 'inline':
            heading = self.doc.add_heading(level=min(level, 6))
            self._process_inline_content(tokens[content_idx], heading)
        
        # 跳過 heading_close
        return start_idx + 3
    
    def _process_paragraph(self, tokens: List[Token], start_idx: int) -> int:
        """處理段落"""
        # 找到段落內容
        content_idx = start_idx + 1
        if content_idx < len(tokens) and tokens[content_idx].type == 'inline':
            para = self.doc.add_paragraph()
            self._process_inline_content(tokens[content_idx], para)
        
        # 跳過 paragraph_close
        return start_idx + 3
    
    def _process_list(self, tokens: List[Token], start_idx: int, ordered: bool = False) -> int:
        """處理列表"""
        i = start_idx + 1
        
        while i < len(tokens):
            token = tokens[i]
            
            if token.type == 'list_item_open':
                i = self._process_list_item(tokens, i, ordered)
            elif token.type in ['bullet_list_close', 'ordered_list_close']:
                return i + 1
            else:
                i += 1
        
        return i
    
    def _process_list_item(self, tokens: List[Token], start_idx: int, ordered: bool = False) -> int:
        """處理列表項"""
        i = start_idx + 1
        
        # 找到列表項的段落內容
        while i < len(tokens) and tokens[i].type != 'list_item_close':
            if tokens[i].type == 'paragraph_open':
                content_idx = i + 1
                if content_idx < len(tokens) and tokens[content_idx].type == 'inline':
                    style = 'List Number' if ordered else 'List Bullet'
                    try:
                        para = self.doc.add_paragraph(style=style)
                        self._process_inline_content(tokens[content_idx], para)
                    except:
                        # 如果樣式不存在，使用普通段落
                        para = self.doc.add_paragraph()
                        marker = "• " if not ordered else ""
                        if marker:
                            run = para.add_run(marker)
                        self._process_inline_content(tokens[content_idx], para)
                i += 3  # 跳過 paragraph_open, inline, paragraph_close
            else:
                i += 1
        
        return i + 1  # 跳過 list_item_close
    
    def _process_blockquote(self, tokens: List[Token], start_idx: int) -> int:
        """處理引用塊"""
        i = start_idx + 1
        
        while i < len(tokens) and tokens[i].type != 'blockquote_close':
            if tokens[i].type == 'paragraph_open':
                content_idx = i + 1
                if content_idx < len(tokens) and tokens[content_idx].type == 'inline':
                    try:
                        para = self.doc.add_paragraph(style='Quote')
                    except:
                        para = self.doc.add_paragraph()
                        para.style = self.doc.styles['Normal']
                    self._process_inline_content(tokens[content_idx], para)
                i += 3
            else:
                i += 1
        
        return i + 1
    
    def _process_code_block(self, token: Token):
        """處理代碼塊"""
        code_text = token.content.rstrip('\n')
        try:
            para = self.doc.add_paragraph(code_text, style='CodeBlock')
        except:
            para = self.doc.add_paragraph(code_text)
            para.style = self.doc.styles['Normal']
    
    def _process_inline_content(self, token: Token, paragraph):
        """處理內聯內容（粗體、斜體等）"""
        if hasattr(token, 'children') and token.children:
            for child in token.children:
                self._process_inline_token(child, paragraph)
        else:
            # 如果沒有子元素，直接添加內容
            if token.content:
                paragraph.add_run(token.content)
    
    def _process_inline_token(self, token: Token, paragraph):
        """處理單個內聯 token"""
        if token.type == 'text':
            paragraph.add_run(token.content)
        elif token.type == 'strong_open':
            # 開始粗體，不添加內容
            pass
        elif token.type == 'strong_close':
            # 結束粗體，不添加內容
            pass
        elif token.type == 'em_open':
            # 開始斜體，不添加內容
            pass
        elif token.type == 'em_close':
            # 結束斜體，不添加內容
            pass
        elif token.type == 'code_inline':
            run = paragraph.add_run(token.content)
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        elif token.type == 's_open':
            # 開始刪除線
            pass
        elif token.type == 's_close':
            # 結束刪除線
            pass
        else:
            # 處理嵌套的格式化內容
            if hasattr(token, 'children') and token.children:
                self._process_nested_formatting(token, paragraph)
    
    def _process_nested_formatting(self, token: Token, paragraph):
        """處理嵌套的格式化內容"""
        # 收集所有文本內容和格式
        content_parts = []
        self._collect_formatted_content(token, content_parts, [])
        
        # 添加格式化的內容到段落
        for text, formats in content_parts:
            if text:
                run = paragraph.add_run(text)
                if 'bold' in formats:
                    run.bold = True
                if 'italic' in formats:
                    run.italic = True
                if 'code' in formats:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                if 'strikethrough' in formats:
                    run.font.strike = True
    
    def _collect_formatted_content(self, token: Token, content_parts: List, current_formats: List):
        """遞歸收集格式化內容"""
        new_formats = current_formats.copy()
        
        # 根據 token 類型添加格式
        if token.type == 'strong_open':
            new_formats.append('bold')
        elif token.type == 'em_open':
            new_formats.append('italic')
        elif token.type == 'code_inline':
            content_parts.append((token.content, new_formats + ['code']))
            return
        elif token.type == 's_open':
            new_formats.append('strikethrough')
        elif token.type == 'text':
            content_parts.append((token.content, new_formats))
            return
        
        # 處理子元素
        if hasattr(token, 'children') and token.children:
            for child in token.children:
                self._collect_formatted_content(child, content_parts, new_formats)
    
    def _add_horizontal_line(self):
        """添加水平線"""
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run("━" * 50)


def convert_markdown_to_docx_improved(markdown_content: str, title: str = "文檔") -> BytesIO:
    """
    改進的便捷函數：將 Markdown 轉換為 DOCX
    
    Args:
        markdown_content (str): Markdown 內容
        title (str): 文檔標題
        
    Returns:
        BytesIO: DOCX 文檔流
    """
    converter = ImprovedMarkdownToDocxConverter()
    return converter.convert_markdown_to_docx(markdown_content, title)