"""
Markdown to DOCX Converter
將 Markdown 格式的內容轉換為 Word DOCX 文檔，保留格式和樣式
"""

import re
import logging
from typing import Dict, List, Optional, Tuple
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

logger = logging.getLogger(__name__)


class MarkdownToDocxConverter:
    """Markdown 轉 DOCX 轉換器"""
    
    def __init__(self):
        self.doc = Document()
        self.setup_styles()
    
    def setup_styles(self):
        """設置 Word 文檔樣式"""
        try:
            # 設置正文樣式
            normal_style = self.doc.styles['Normal']
            normal_font = normal_style.font
            normal_font.name = '微軟正黑體'
            normal_font.size = Pt(11)
            
            # 創建代碼塊樣式
            try:
                code_style = self.doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
                code_font = code_style.font
                code_font.name = 'Consolas'
                code_font.size = Pt(10)
                code_style.paragraph_format.left_indent = Inches(0.5)
                code_style.paragraph_format.space_before = Pt(6)
                code_style.paragraph_format.space_after = Pt(6)
            except:
                pass  # 樣式可能已存在
            
            # 創建引用樣式
            try:
                quote_style = self.doc.styles.add_style('Quote', WD_STYLE_TYPE.PARAGRAPH)
                quote_style.paragraph_format.left_indent = Inches(0.5)
                quote_style.paragraph_format.space_before = Pt(6)
                quote_style.paragraph_format.space_after = Pt(6)
                quote_font = quote_style.font
                quote_font.italic = True
                quote_font.color.rgb = None  # 灰色
            except:
                pass
                
        except Exception as e:
            logger.warning(f"設置樣式時發生錯誤: {e}")
    
    def convert_markdown_to_docx(self, markdown_content: str, title: str = "文檔") -> BytesIO:
        """
        將 Markdown 內容轉換為 DOCX 文檔
        
        Args:
            markdown_content (str): Markdown 格式的內容
            title (str): 文檔標題
            
        Returns:
            BytesIO: DOCX 文檔的二進制流
        """
        try:
            # 添加文檔標題
            if title and title != "文檔":
                title_para = self.doc.add_heading(title, level=0)
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 分行處理
            lines = markdown_content.split('\n')
            i = 0
            
            while i < len(lines):
                line = lines[i].rstrip()
                
                if not line:  # 空行
                    self.doc.add_paragraph()
                    i += 1
                    continue
                
                # 處理代碼塊
                if line.startswith('```'):
                    i = self._process_code_block(lines, i)
                    continue
                
                # 處理標題
                if line.startswith('#'):
                    self._process_heading(line)
                    i += 1
                    continue
                
                # 處理引用
                if line.startswith('>'):
                    i = self._process_blockquote(lines, i)
                    continue
                
                # 處理列表
                if self._is_list_item(line):
                    i = self._process_list(lines, i)
                    continue
                
                # 處理水平線
                if line.strip() in ['---', '***', '___']:
                    self._add_horizontal_line()
                    i += 1
                    continue
                
                # 處理普通段落
                self._process_paragraph(line)
                i += 1
            
            # 生成 DOCX 二進制流
            docx_stream = BytesIO()
            self.doc.save(docx_stream)
            docx_stream.seek(0)
            
            return docx_stream
            
        except Exception as e:
            logger.error(f"轉換 Markdown 到 DOCX 時發生錯誤: {e}")
            raise
    
    def _process_heading(self, line: str):
        """處理標題"""
        level = 0
        for char in line:
            if char == '#':
                level += 1
            else:
                break
        
        heading_text = line[level:].strip()
        if heading_text:
            # 處理標題中的內聯格式 (如 ## **粗體標題**)
            heading = self.doc.add_heading(level=min(level, 6))
            self._add_formatted_text(heading, heading_text)
    
    def _process_code_block(self, lines: List[str], start_idx: int) -> int:
        """處理代碼塊"""
        i = start_idx + 1
        code_lines = []
        
        # 收集代碼塊內容
        while i < len(lines) and not lines[i].rstrip().startswith('```'):
            code_lines.append(lines[i])
            i += 1
        
        # 添加代碼塊
        if code_lines:
            code_text = '\n'.join(code_lines)
            try:
                para = self.doc.add_paragraph(code_text, style='CodeBlock')
            except:
                # 如果樣式不存在，使用普通段落
                para = self.doc.add_paragraph(code_text)
                para.style = self.doc.styles['Normal']
        
        return i + 1 if i < len(lines) else i
    
    def _process_blockquote(self, lines: List[str], start_idx: int) -> int:
        """處理引用塊"""
        i = start_idx
        quote_lines = []
        
        # 收集引用內容
        while i < len(lines) and lines[i].strip().startswith('>'):
            quote_text = lines[i].strip()[1:].strip()
            if quote_text:
                quote_lines.append(quote_text)
            i += 1
        
        # 添加引用
        if quote_lines:
            quote_text = '\n'.join(quote_lines)
            try:
                para = self.doc.add_paragraph(quote_text, style='Quote')
            except:
                para = self.doc.add_paragraph(quote_text)
                para.style = self.doc.styles['Normal']
        
        return i
    
    def _is_list_item(self, line: str) -> bool:
        """檢查是否為列表項"""
        stripped = line.strip()
        # 無序列表
        if stripped.startswith(('- ', '* ', '+ ')):
            return True
        # 有序列表
        if re.match(r'^\d+\.\s', stripped):
            return True
        return False
    
    def _process_list(self, lines: List[str], start_idx: int) -> int:
        """處理列表"""
        i = start_idx
        
        while i < len(lines) and self._is_list_item(lines[i]):
            line = lines[i].strip()
            
            # 移除列表標記
            if line.startswith(('- ', '* ', '+ ')):
                list_text = line[2:].strip()
                para = self.doc.add_paragraph(list_text, style='List Bullet')
            elif re.match(r'^\d+\.\s', line):
                list_text = re.sub(r'^\d+\.\s', '', line).strip()
                para = self.doc.add_paragraph(list_text, style='List Number')
            
            i += 1
        
        return i
    
    def _process_paragraph(self, line: str):
        """處理普通段落，支持內聯格式"""
        paragraph = self.doc.add_paragraph()
        
        # 處理內聯格式
        self._add_formatted_text(paragraph, line)
    
    def _add_formatted_text(self, paragraph, text: str):
        """添加格式化文本到段落"""
        # 處理粗體、斜體、行內代碼等
        parts = self._parse_inline_formatting(text)
        
        for part_text, formatting in parts:
            run = paragraph.add_run(part_text)
            
            if 'bold' in formatting:
                run.bold = True
            if 'italic' in formatting:
                run.italic = True
            if 'code' in formatting:
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
            if 'strikethrough' in formatting:
                run.font.strike = True
    
    def _parse_inline_formatting(self, text: str) -> List[Tuple[str, List[str]]]:
        """解析內聯格式 - 使用正則表達式進行更可靠的解析"""
        import re
        
        parts = []
        current_pos = 0
        
        # 定義所有格式的正則表達式模式 (順序很重要：先處理較長的模式)
        patterns = [
            (r'\*\*(.+?)\*\*', 'bold'),      # **bold** (先處理雙星號)
            (r'(?<!\*)\*([^*]+?)\*(?!\*)', 'italic'),  # *italic* (但不是 **text**)
            (r'`(.+?)`', 'code'),            # `code`
            (r'~~(.+?)~~', 'strikethrough'), # ~~strikethrough~~
        ]
        
        # 找到所有匹配項及其位置
        matches = []
        for pattern, format_type in patterns:
            for match in re.finditer(pattern, text):
                matches.append((match.start(), match.end(), match.group(1), format_type))
        
        # 按位置排序
        matches.sort(key=lambda x: x[0])
        
        # 處理重疊的匹配項（優先處理較長的匹配）
        filtered_matches = []
        for match in matches:
            start, end, content, format_type = match
            # 檢查是否與已有匹配項重疊
            overlaps = False
            for existing in filtered_matches:
                existing_start, existing_end = existing[0], existing[1]
                if not (end <= existing_start or start >= existing_end):
                    overlaps = True
                    break
            
            if not overlaps:
                filtered_matches.append(match)
        
        # 重新按位置排序
        filtered_matches.sort(key=lambda x: x[0])
        
        # 構建結果
        for match in filtered_matches:
            start, end, content, format_type = match
            
            # 添加匹配項之前的普通文本
            if start > current_pos:
                plain_text = text[current_pos:start]
                if plain_text:
                    parts.append((plain_text, []))
            
            # 添加格式化文本
            parts.append((content, [format_type]))
            current_pos = end
        
        # 添加最後的普通文本
        if current_pos < len(text):
            remaining_text = text[current_pos:]
            if remaining_text:
                parts.append((remaining_text, []))
        
        # 如果沒有找到任何格式，返回整個文本作為普通文本
        if not parts:
            parts.append((text, []))
        
        return parts
    
    def _add_horizontal_line(self):
        """添加水平線"""
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run("━" * 50)
        run.font.color.rgb = None  # 設為灰色


def convert_markdown_to_docx(markdown_content: str, title: str = "文檔") -> BytesIO:
    """
    便捷函數：將 Markdown 轉換為 DOCX
    
    Args:
        markdown_content (str): Markdown 內容
        title (str): 文檔標題
        
    Returns:
        BytesIO: DOCX 文檔流
    """
    converter = MarkdownToDocxConverter()
    return converter.convert_markdown_to_docx(markdown_content, title)